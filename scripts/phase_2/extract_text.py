import os
import time
from typing import Any, List, Tuple, Optional
import psycopg2
from psycopg2 import OperationalError
from PyPDF2 import PdfReader
from docx import Document

from scripts.config.phase_2 import TEXT_CHAR_LIMIT, LOG_FILE

def log(msg: str) -> None:
    """Registra un mensaje en el log central y lo emite por consola."""
    os.makedirs(os.path.dirname(LOG_FILE), exist_ok=True)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(msg + "\n")
    print(msg)

def get_db_connection(retries: int = 10, delay: int = 3) -> Any:
    """
    Establece conexión con PostgreSQL con reintentos y backoff exponencial.
    """
    for attempt in range(1, retries + 1):
        try:
            return psycopg2.connect(
                dbname=os.getenv("PGDATABASE", os.getenv("POSTGRES_DB", "auditdb")),
                user=os.getenv("PGUSER", os.getenv("POSTGRES_USER", "user")),
                password=os.getenv("PGPASSWORD", os.getenv("POSTGRES_PASSWORD", "pass")),
                host=os.getenv("PGHOST"),
                port=int(os.getenv("PGPORT", "5432")),
            )
        except OperationalError as e:
            log(f"Postgres not ready (attempt {attempt}/{retries}): {e}")
            time.sleep(delay)
    raise RuntimeError("Could not connect to Postgres.")

def safe_read_text_file(path: str, max_chars: int) -> str:
    """
    Lee archivos de texto plano probando múltiples codificaciones.
    
    Intenta decodificar en orden: UTF-8, Latin-1 y CP1252 para maximizar 
    la compatibilidad con archivos antiguos o de diferentes sistemas operativos.
    """
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=enc, errors="ignore") as f:
                content = f.read(max_chars + 1000)
                return content[:max_chars]
        except Exception:
            continue
    return ""

def extract_text_from_pdf(path: str, max_chars: int, ocr_needed: bool) -> str:
    """
    Extrae la capa de texto digital de un PDF.
    
    Args:
        path: Ruta al archivo.
        max_chars: Límite de caracteres a extraer.
        ocr_needed: Si es True, el script ignora el archivo para procesarlo 
                    en una fase posterior con modelos OCR.
    """
    if ocr_needed:
        return ""

    try:
        reader = PdfReader(path)
        parts: List[str] = []
        total_len = 0

        for page in reader.pages:
            text = page.extract_text() or ""
            if text:
                parts.append(text)
                total_len += len(text)
            if total_len >= max_chars:
                break

        return "\n".join(parts)[:max_chars]
    except Exception as e:
        log(f"Error reading PDF {path}: {e}")
        return ""

def extract_text_from_docx(path: str, max_chars: int) -> str:
    """
    Extrae texto de archivos Word (.docx) procesando párrafos y tablas.
    """
    try:
        doc = Document(path)
        parts: List[str] = []
        total_len = 0

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
                total_len += len(p.text)
                if total_len >= max_chars: break

        if total_len < max_chars:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
                            total_len += len(cell.text)
                            if total_len >= max_chars: break
        
        return "\n".join(parts)[:max_chars]
    except Exception as e:
        log(f"Error reading DOCX {path}: {e}")
        return ""

def main(batch_size: int = 500) -> None:
    """
    Orquesta la extracción masiva de texto en lotes (batch processing).
    
    El flujo identifica archivos sin `text_excerpt`, determina su tipo 
    y utiliza el extractor correspondiente. Los resultados se guardan 
    en la base de datos para habilitar la búsqueda semántica.
    """
    conn = get_db_connection()
    cur = conn.cursor()

    log("Starting text extraction for Phase 2...")
    total_processed = 0

    while True:
        cur.execute(
            "SELECT id FROM files WHERE text_excerpt IS NULL ORDER BY id LIMIT %s",
            (batch_size,)
        )
        ids = [row[0] for row in cur.fetchall()]
        if not ids: break

        updates: List[Tuple[str, int, int]] = []

        for file_id in ids:
            cur.execute(
                "SELECT full_path, file_type, is_pdf, ocr_needed FROM files WHERE id = %s",
                (file_id,)
            )
            row = cur.fetchone()
            if not row or not row[0] or not os.path.exists(row[0]):
                updates.append(("", 0, file_id))
                continue

            full_path, file_type, is_pdf, ocr_needed = row
            text = ""

            try:
                if is_pdf:
                    text = extract_text_from_pdf(full_path, TEXT_CHAR_LIMIT, bool(ocr_needed))
                else:
                    ext = (file_type or "").lower()
                    if ext in (".txt", ".log", ".md"):
                        text = safe_read_text_file(full_path, TEXT_CHAR_LIMIT)
                    elif ext == ".docx":
                        text = extract_text_from_docx(full_path, TEXT_CHAR_LIMIT)
            except Exception as e:
                log(f"Error in id={file_id}: {e}")
            
            updates.append((text, len(text), file_id))
            total_processed += 1

        if updates:
            cur.executemany(
                "UPDATE files SET text_excerpt = %s, text_chars_extracted = %s WHERE id = %s",
                updates
            )
            conn.commit()
            log(f"Processed {total_processed} files...")

    cur.close()
    conn.close()
    log("=== Text extraction completed ===")

if __name__ == "__main__":
    main()